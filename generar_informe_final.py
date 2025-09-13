import os
from datetime import datetime, date, timedelta

try:
    from docx import Document
    from docx.shared import Pt, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except Exception as e:
    raise SystemExit("Falta dependencia 'python-docx'. Instálala con: pip install python-docx")


def add_heading(document: Document, text: str, level: int = 1):
    paragraph = document.add_heading(text, level=level)
    return paragraph


def add_paragraph(document: Document, text: str, bold: bool = False, italic: bool = False, align_center: bool = False):
    paragraph = document.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = bold
    run.italic = italic
    paragraph_format = paragraph.paragraph_format
    if align_center:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    font = run.font
    font.name = 'Calibri'
    font.size = Pt(11)
    return paragraph


def add_numbered_list(document: Document, items):
    for item in items:
        p = document.add_paragraph(style='List Number')
        r = p.add_run(item)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def add_bulleted_list(document: Document, items):
    for item in items:
        p = document.add_paragraph(style='List Bullet')
        r = p.add_run(item)
        r.font.name = 'Calibri'
        r.font.size = Pt(11)


def add_page_break(document: Document):
    document.add_page_break()


def add_footer_page_numbers(document: Document):
    section = document.sections[0]
    footer = section.footer
    paragraph = footer.paragraphs[0]
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Add PAGE field
    run = paragraph.add_run()
    fldChar1 = OxmlElement('w:fldChar')
    fldChar1.set(qn('w:fldCharType'), 'begin')
    instrText = OxmlElement('w:instrText')
    instrText.text = 'PAGE'
    fldChar2 = OxmlElement('w:fldChar')
    fldChar2.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar1)
    run._r.append(instrText)
    run._r.append(fldChar2)


def get_weeks_jan_to_sep_2025():
    year = 2025
    start = date(year, 1, 1)
    # Primer lunes del año
    delta_to_monday = (0 - start.weekday()) % 7
    first_monday = start + timedelta(days=delta_to_monday)
    end = date(year, 9, 30)

    weeks = []
    current = first_monday
    while current <= end:
        week_start = current
        week_end = min(current + timedelta(days=6), end)
        weeks.append({
            'index': len(weeks) + 1,
            'start': week_start,
            'end': week_end
        })
        current += timedelta(days=7)
    return weeks


def build_planned_tasks():
    # start_week es 1-based respecto a get_weeks_jan_to_sep_2025()
    return [
        { 'name': 'Levantamiento de requisitos, alcance y riesgos', 'module': 'Arquitectura/Planificación', 'start_week': 1, 'duration': 2, 'details': 'Entrevistas, definición de módulos, criterios de aceptación.' },
        { 'name': 'Arquitectura y configuración base', 'module': 'Arquitectura/Setup', 'start_week': 3, 'duration': 1, 'details': 'App factory, blueprints, configuración.' },
        { 'name': 'Modelado BD inicial (Usuario, Paciente, Clinica)', 'module': 'Base de Datos', 'start_week': 4, 'duration': 1, 'details': 'SQLAlchemy, migraciones, relaciones base.' },
        { 'name': 'Autenticación básica (login/logout)', 'module': 'Auth', 'start_week': 5, 'duration': 1, 'details': 'Rutas, formularios, templates, sesión.' },
        { 'name': 'Registro/gestión de usuarios y roles', 'module': 'Auth/Admin', 'start_week': 6, 'duration': 1, 'details': 'Registro, cambio de rol, vistas admin.' },
        { 'name': 'Módulo Pacientes: CRUD, formularios, templates', 'module': 'Pacientes', 'start_week': 7, 'duration': 1, 'details': 'Alta/edición/borrado, validación, listado.' },
        { 'name': 'Búsqueda avanzada de pacientes', 'module': 'Pacientes', 'start_week': 8, 'duration': 1, 'details': 'Filtros, AJAX, optimizaciones de consulta.' },
        { 'name': 'Estructura módulo Consultas', 'module': 'Consultas', 'start_week': 9, 'duration': 1, 'details': 'Rutas, templates, tabs de datos.' },
        { 'name': 'Signos vitales y modelo SignosVitales', 'module': 'Consultas', 'start_week': 10, 'duration': 1, 'details': 'Modelo, formulario, persistencia.' },
        { 'name': 'Motivo de consulta e historia clínica', 'module': 'Consultas', 'start_week': 11, 'duration': 1, 'details': 'Anamnesis y revisión por sistemas.' },
        { 'name': 'Diagnóstico, tratamiento y receta', 'module': 'Consultas', 'start_week': 12, 'duration': 1, 'details': 'Persistencia, presentación e impresión.' },
        { 'name': 'Autoguardado y mejoras JS', 'module': 'Frontend/JS', 'start_week': 13, 'duration': 1, 'details': 'LocalStorage, intervalos, indicadores.' },
        { 'name': 'Asignación de clínicas a médicos', 'module': 'Admin', 'start_week': 14, 'duration': 1, 'details': 'Vistas de asignación y control de acceso.' },
        { 'name': 'Auditoría de cambios (roles/asignaciones)', 'module': 'Auditoría', 'start_week': 15, 'duration': 1, 'details': 'Historial de cambios, responsable y timestamp.' },
        { 'name': 'Reportes diarios y por clínica', 'module': 'Reportes', 'start_week': 16, 'duration': 1, 'details': 'Consultas agregadas y filtros.' },
        { 'name': 'Reportes de pacientes y estadísticas', 'module': 'Reportes/Gráficos', 'start_week': 17, 'duration': 1, 'details': 'Gráficas con plot_generator.' },
        { 'name': 'Generación de PDFs (informe médico, reportes)', 'module': 'Reportes/PDF', 'start_week': 18, 'duration': 1, 'details': 'ReportLab, plantillas y estilos.' },
        { 'name': 'Seguridad y validaciones', 'module': 'Seguridad', 'start_week': 19, 'duration': 1, 'details': 'CSRF, WTForms, hashing de contraseñas.' },
        { 'name': 'Correo (reset password) y configuración', 'module': 'Email', 'start_week': 20, 'duration': 1, 'details': 'Flask-Mail, plantillas de correo.' },
        { 'name': 'Backups y utilidades', 'module': 'Operaciones', 'start_week': 21, 'duration': 1, 'details': 'Scripts de backup y pruebas con Drive.' },
        { 'name': 'UI/UX y responsive con Bootstrap', 'module': 'Frontend/UI', 'start_week': 22, 'duration': 1, 'details': 'Refinamiento visual y accesibilidad.' },
        { 'name': 'Pruebas integrales y correcciones', 'module': 'QA', 'start_week': 23, 'duration': 1, 'details': 'Flujo completo y errores críticos.' },
        { 'name': 'Configuración MySQL y migraciones', 'module': 'Infraestructura', 'start_week': 24, 'duration': 1, 'details': 'DSNs, ajustes ORM, pruebas conexión.' },
        { 'name': 'Dashboards y métricas', 'module': 'Reportes/Estadísticas', 'start_week': 25, 'duration': 1, 'details': 'Indicadores clave en dashboard.' },
        { 'name': 'Historial clínico y navegación', 'module': 'Consultas', 'start_week': 26, 'duration': 1, 'details': 'Listado histórico por paciente.' },
        { 'name': 'Estados de consulta y flujos', 'module': 'Consultas', 'start_week': 27, 'duration': 1, 'details': 'en_progreso/completada y transiciones.' },
        { 'name': 'Guías de despliegue (Gunicorn/Nginx)', 'module': 'DevOps', 'start_week': 28, 'duration': 1, 'details': 'config, systemd, estáticos.' },
        { 'name': 'Mejoras de seguridad adicionales', 'module': 'Seguridad', 'start_week': 29, 'duration': 1, 'details': 'Timeout sesión, sanitización extra.' },
        { 'name': 'Refactor de templates y forms', 'module': 'Limpieza Técnica', 'start_week': 30, 'duration': 1, 'details': 'Reuso de componentes y DRY.' },
        { 'name': 'Datos ficticios para pruebas', 'module': 'Datos/QA', 'start_week': 31, 'duration': 1, 'details': 'Scripts de carga y volúmenes.' },
        { 'name': 'Campos gineco-obstétricos condicionales', 'module': 'Consultas', 'start_week': 32, 'duration': 1, 'details': 'Visibilidad y validación por género.' },
        { 'name': 'Rendimiento JS y búsqueda con debounce', 'module': 'Frontend/JS', 'start_week': 33, 'duration': 1, 'details': 'Optimización de UI y llamadas.' },
        { 'name': 'Piloto interno y ajustes', 'module': 'Piloto', 'start_week': 34, 'duration': 1, 'details': 'Feedback de usuarios y fixes.' },
        { 'name': 'Presentación ejecutiva y guía de uso', 'module': 'Documentación', 'start_week': 35, 'duration': 1, 'details': 'Material de exposición y manuales.' },
        { 'name': 'Documentación técnica completa', 'module': 'Documentación', 'start_week': 36, 'duration': 1, 'details': 'Arquitectura, API, despliegue.' },
        { 'name': 'Piloto ampliado y soporte', 'module': 'Piloto/Soporte', 'start_week': 37, 'duration': 2, 'details': 'Soporte en sitio y monitoreo.' },
        { 'name': 'Cierre, informe final y anexos', 'module': 'Cierre', 'start_week': 39, 'duration': 1, 'details': 'Informe final, anexos, lecciones.' },
    ]


def add_schedule_table(document: Document):
    weeks = get_weeks_jan_to_sep_2025()
    tasks = build_planned_tasks()

    # Mapa semana -> lista de tareas (nombre, módulo)
    week_to_items = { w['index']: [] for w in weeks }
    for t in tasks:
        for w in range(t['start_week'], t['start_week'] + t['duration']):
            if w in week_to_items:
                week_to_items[w].append((t['name'], t['module']))

    add_paragraph(document, 'Cronograma detallado por semanas (enero a septiembre 2025). Cada fila representa una semana calendario (lunes a domingo).')

    table = document.add_table(rows=1, cols=5)
    try:
        table.style = 'Light List Accent 1'
    except Exception:
        pass
    hdr_cells = table.rows[0].cells
    hdr_cells[0].text = 'Semana'
    hdr_cells[1].text = 'Inicio'
    hdr_cells[2].text = 'Fin'
    hdr_cells[3].text = 'Actividades/Entregables'
    hdr_cells[4].text = 'Módulos'

    for w in weeks:
        items = week_to_items[w['index']]
        actividades = '\n'.join(f"- {name}" for name, _ in items) if items else '- Documentación y soporte'
        modulos = ', '.join(sorted(set(m for _, m in items))) if items else 'Documentación'

        row_cells = table.add_row().cells
        row_cells[0].text = f"W{w['index']}"
        row_cells[1].text = w['start'].strftime('%d/%m/%Y')
        row_cells[2].text = w['end'].strftime('%d/%m/%Y')
        row_cells[3].text = actividades
        row_cells[4].text = modulos

def build_document():
    document = Document()

    # Configuración de márgenes
    for section in document.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1)
        section.right_margin = Inches(1)

    # Portada
    add_heading(document, 'INFORME FINAL', level=0)
    add_paragraph(document, 'Sistema de Gestión Clínica CUNORI - Shororagua', bold=True, align_center=True)
    add_paragraph(document, f'Fecha: {datetime.now().strftime("%d/%m/%Y")}', align_center=True)
    add_page_break(document)

    # Índice (placeholder; Word puede actualizar índices automáticamente)
    add_heading(document, 'Contenido', level=1)
    add_bulleted_list(document, [
        'Resumen del proyecto',
        'Introducción',
        'Justificación',
        'Alcance',
        '1. Vista General del Proyecto',
        '  1.1 Objetivos',
        '  1.2 Metodología',
        '  1.3 Recursos',
        '  1.4 Entregables del proyecto',
        '2. Organización del proyecto',
        '  2.1 Participantes en el proyecto',
        '  2.2 Roles y Responsabilidades',
        '  2.3 Cronograma de actividades',
        '3. Desarrollo del Proyecto',
        'Conclusiones',
        'Recomendaciones',
        'Referencias Bibliográficas',
        'E-Grafía',
        'Apéndice',
        'Glosario',
        'Anexos',
        'Encabezado 2',
    ])
    add_page_break(document)

    # Secciones
    add_heading(document, 'Resumen del proyecto', level=1)
    add_paragraph(document, (
        'El Sistema de Gestión Clínica moderniza la atención médica y la administración de clínicas '
        'universitarias mediante digitalización de expedientes, control de consultas, gestión de roles '
        'y generación de reportes. Implementado con Flask, SQLAlchemy y una arquitectura MVC, el '
        'sistema habilita trazabilidad completa, seguridad y escalabilidad, reduciendo tiempos operativos '
        'y mejorando la calidad de atención.'
    ))

    add_heading(document, 'Introducción', level=1)
    add_paragraph(document, (
        'Este informe documenta de forma integral el desarrollo del sistema, sus objetivos, alcance, '
        'arquitectura, módulos, organización del equipo, cronograma y resultados, proporcionando una '
        'visión holística del proyecto y su impacto.'
    ))

    add_heading(document, 'Justificación', level=1)
    add_paragraph(document, (
        'La gestión manual con expedientes físicos implica ineficiencias, errores y falta de control. '
        'La solución digital propuesta resuelve estas limitaciones mediante procesos estandarizados, '
        'datos centralizados y reportes automáticos, favoreciendo decisiones basadas en datos.'
    ))

    add_heading(document, 'Alcance', level=1)
    add_bulleted_list(document, [
        'Gestión de pacientes (registro, actualización, historial).',
        'Gestión de consultas (motivo, diagnóstico, tratamiento, receta).',
        'Signos vitales y examen físico.',
        'Módulo administrativo con roles y permisos.',
        'Asignación de clínicas a médicos.',
        'Reportes diarios, por clínica, pacientes y estadísticos.',
        'Auditoría de cambios de roles y asignaciones.'
    ])

    add_heading(document, '1. Vista General del Proyecto', level=1)

    add_heading(document, '1.1 Objetivos', level=2)
    add_bulleted_list(document, [
        'Digitalizar procesos clínicos y administrativos.',
        'Mejorar eficiencia y reducir errores.',
        'Asegurar trazabilidad y auditoría completa.',
        'Proveer reportes y métricas en tiempo real.',
        'Escalar a múltiples clínicas y usuarios.'
    ])

    add_heading(document, '1.2 Metodología', level=2)
    add_paragraph(document, (
        'Se adoptó una metodología iterativa e incremental, con ciclos de análisis, diseño, '
        'implementación y validación, entregas parciales demostrables y retroalimentación continua '
        'de usuarios clave.'
    ))

    add_heading(document, '1.3 Recursos', level=2)
    add_bulleted_list(document, [
        'Equipo: administración, médicos supervisores, médicos, recepción, desarrollador(es).',
        'Tecnología: Python 3.10+, Flask, SQLAlchemy, Bootstrap, SQLite/MySQL.',
        'Infraestructura: servidor de aplicación, base de datos, almacenamiento para backups.'
    ])

    add_heading(document, '1.4 Entregables del proyecto', level=2)
    add_bulleted_list(document, [
        'Código fuente y estructura modular por blueprints.',
        'Base de datos con modelos y migraciones.',
        'Reportes PDF y gráficos estadísticos.',
        'Documentación técnica y guía de uso.',
        'Presentación ejecutiva y material de capacitación.'
    ])

    add_heading(document, '2. Organización del proyecto', level=1)

    add_heading(document, '2.1 Participantes en el proyecto', level=2)
    add_bulleted_list(document, [
        'Patrocinador académico',
        'Líder técnico / Desarrollador',
        'Médico supervisor (experto de dominio)',
        'Médicos usuarios',
        'Recepción y personal administrativo'
    ])

    add_heading(document, '2.2 Roles y Responsabilidades', level=2)
    add_bulleted_list(document, [
        'Administrador: configuración, usuarios, roles y clínicas.',
        'Médico supervisor: gestión de médicos y reportes avanzados.',
        'Médico: consultas, diagnóstico, receta e historial.',
        'Recepción: registro y actualización de pacientes.'
    ])

    add_heading(document, '2.3 Cronograma de actividades', level=2)
    add_schedule_table(document)

    add_heading(document, '3. Desarrollo del Proyecto', level=1)
    add_paragraph(document, (
        'El desarrollo se estructuró por módulos (auth, main, pacientes, reportes) bajo blueprints. '
        'Se implementaron modelos `Usuario`, `Paciente`, `Consulta`, `SignosVitales` y `Clinica`. '
        'La autenticación usa Flask-Login con roles. Se generaron reportes en PDF y gráficos. '
        'Se adoptaron validaciones de formularios, control de acceso por rol y auditoría de cambios.'
    ))

    add_heading(document, 'Conclusiones', level=1)
    add_bulleted_list(document, [
        'El sistema moderniza y estandariza procesos clínicos y administrativos.',
        'La trazabilidad y seguridad fortalecen el control institucional.',
        'La arquitectura modular facilita mantenimiento y futuras extensiones.'
    ])

    add_heading(document, 'Recomendaciones', level=1)
    add_bulleted_list(document, [
        'Implementar autenticación multifactor para perfiles críticos.',
        'Desplegar monitoreo de desempeño y registros centralizados.',
        'Plan de continuidad operativa y políticas de respaldo/retención.',
        'Futuras integraciones con sistemas externos mediante APIs.'
    ])

    add_heading(document, 'Referencias Bibliográficas', level=1)
    add_bulleted_list(document, [
        'Documentación de Flask y extensiones asociadas.',
        'Buenas prácticas de seguridad OWASP para aplicaciones web.',
        'Normativas de manejo de información en el contexto institucional.'
    ])

    add_heading(document, 'E-Grafía', level=1)
    add_bulleted_list(document, [
        'Repositorio del proyecto y documentación técnica adjunta.',
        'Guías de uso y presentaciones ejecutivas incluidas en el proyecto.'
    ])

    add_heading(document, 'Apéndice', level=1)
    add_bulleted_list(document, [
        'Capturas de pantallas de módulos principales.',
        'Ejemplos de reportes y PDFs generados.'
    ])

    add_heading(document, 'Glosario', level=1)
    add_bulleted_list(document, [
        'Consulta: Interacción clínica registrada con un paciente.',
        'Expediente: Conjunto de datos médicos del paciente.',
        'Rol: Perfil de permisos asignado a un usuario.'
    ])

    add_heading(document, 'Anexos', level=1)
    add_bulleted_list(document, [
        'Diagramas de arquitectura y base de datos.',
        'Manual de usuario y material de capacitación.'
    ])

    add_heading(document, 'Encabezado 2', level=1)
    add_paragraph(document, 'Sección adicional solicitada para cumplir con el índice proporcionado.')

    add_footer_page_numbers(document)

    return document


def main():
    output_path = os.path.join(os.path.dirname(__file__), 'INFORME_FINAL.docx')
    document = build_document()
    document.save(output_path)
    print(f"Documento generado correctamente: {output_path}")


if __name__ == '__main__':
    main()



